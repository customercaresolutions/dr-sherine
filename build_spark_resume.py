"""Generate spark-resume.docx — 2-page resume tailored for ocular gene therapy & inherited retinal disorder companies (e.g. Spark Therapeutics), with links to the supporting site."""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

BASE_URL = "https://customercaresolutions.github.io/dr-sherine/versions/spark-therapeutics"


def add_hyperlink(paragraph, url, text, bold=False):
    """Append a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


doc = Document("resume.docx")  # clone margins/styles from existing resume

# Clear all existing content
for el in list(doc.element.body):
    if el.tag == qn("w:p") or el.tag == qn("w:tbl"):
        doc.element.body.remove(el)

# Drop stale hyperlink relationships inherited from the template
HYPERLINK_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
for rel_id in [r for r, rel in doc.part.rels.items() if rel.reltype == HYPERLINK_TYPE]:
    del doc.part.rels[rel_id]

# === HEADER ===
p = doc.add_paragraph()
run = p.add_run("DR. SHERINE MARINA D'SOUZA BRAGANZA")
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph(
    "Vitreoretinal Surgeon | Inherited Retinal Disorders | Ocular Gene Therapy Readiness | Medical Retina"
)

p = doc.add_paragraph(
    "drsherine@gmail.com | +91 98453 66008 | MS Ophthalmology (Gold Medal) | Bangalore, India"
)

p = doc.add_paragraph()
add_hyperlink(p, f"{BASE_URL}/", "Full Online Resume (Ocular Gene Therapy & IRD version)", bold=True)

# === PROFESSIONAL SUMMARY ===
doc.add_heading("PROFESSIONAL SUMMARY", level=2)
doc.add_paragraph(
    "Senior vitreoretinal surgeon and clinical strategist with 25+ years across medical retina, "
    "inherited retinal disorders (IRDs), ROP, and tertiary retinal care. Published on genotype–"
    "phenotype correlation in IRDs and on next-generation sequencing / genetic counseling in ocular "
    "disorders. Helped scale a tele-retina screening network across 50+ centers with 10,000+ patient "
    "records — the same infrastructure used to identify, refer, and follow patients eligible for "
    "advanced retinal therapies. Active physician-educator (16+ supervised DNB theses, 120+ "
    "conference talks, DNB / NBE examiner) with a deep imaging-interpretation footprint (OCT, OCTA, "
    "FAF, FA/ICG, B-scan). The bridge between cutting-edge retinal science and everyday ophthalmic "
    "practice."
)

# === CORE EXPERTISE ===
doc.add_heading("CORE EXPERTISE", level=2)
for item in [
    "Inherited Retinal Disorders (IRDs) • Genotype–Phenotype Correlation • Genetic Counseling",
    "Retinal Imaging Interpretation • OCT / SD-OCT / SS-OCT • OCT-Angiography • FAF • FA/ICG • B-scan",
    "Patient Identification Pathways • Referral Workflows • Disease Progression Monitoring",
    "Treatment-Center Engagement • Multi-Center Network Activation • Tele-Retina Operations",
    "Medical & Surgical Retina • ROP & Pediatric Retina • DR, RVO, AMD, Macular Edema",
    "KOL Engagement • Clinical Advisory • Physician Education (DNB / NBE examiner)",
]:
    doc.add_paragraph(item, style="List Bullet")

# === WHY THIS PROFILE FITS ===
doc.add_heading("ALIGNMENT WITH OCULAR GENE THERAPY & IRD PROGRAMS", level=2)
doc.add_paragraph(
    "Ocular gene therapy programs (companies like Spark Therapeutics) sit at the intersection of "
    "retinal disease, advanced therapeutics, inherited retinal disorders, physician education, and "
    "real-world clinical implementation. These roles call for senior retina specialists who "
    "understand patient identification pathways, retinal imaging interpretation, disease progression, "
    "referral workflows, treatment-center engagement, and how advanced therapies translate into "
    "actual ophthalmic practice. 25+ years in medical retina, ROP, tertiary retinal care, "
    "multidisciplinary teaching, and IRD / genetic-counseling research map directly onto those needs."
)

# === PROFESSIONAL EXPERIENCE ===
doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)

doc.add_paragraph(
    "Narayana Nethralaya II — Senior Consultant & Head, Vitreoretinal Services  (Nov 2006 – Present)\n"
    "Tertiary medical and surgical retina in a high-volume superspeciality eye hospital. Workup and "
    "longitudinal management of inherited retinal disorders, retinal vascular and degenerative "
    "disease, ROP, and complex vitreoretinal pathology. Head of VR Services at the NN2 (Health City) "
    "branch. Lead point for international patients in retinal care. Active in the IRD and "
    "genetic-counseling clinical pathway alongside geneticists and paediatricians. Training fellows "
    "and DNB postgraduates."
)

doc.add_paragraph(
    "KIDROP Initiative — Senior Clinical Lead  (2008 – Present)\n"
    "Helped scale a tele-retinal screening network across 50+ primary and secondary centers, with "
    "cumulative coverage of 10,000+ patient records. Designed clinical workflows, referral pathways, "
    "imaging protocols, and treatment escalation criteria — the same operational building blocks "
    "that underpin a gene-therapy treatment-center network."
)

doc.add_paragraph(
    "Lion's Eye Hospital — Consultant VR Surgeon  (Dec 2003 – Nov 2006)\n"
    "Retinal detachment surgeries, vitrectomy (including 23g), cryopexy, medical retina and lasers, "
    "fluorescein/ICG angiography, B-scan ultrasonography, and OCT."
)

doc.add_paragraph(
    "Prathima Institute of Medical Sciences — Assistant Professor  (Feb 2002 – Present)\n"
    "Undergraduate and postgraduate teaching within a tertiary care teaching hospital (MCI "
    "recognized); supervision of research projects and theses."
)

p = doc.add_paragraph()
add_hyperlink(p, f"{BASE_URL}/full-resume.html", "View full professional experience")

# === EDUCATION & CREDENTIALS ===
doc.add_heading("EDUCATION & CREDENTIALS", level=2)
for item in [
    "MS Ophthalmology — Gold Medal (Rajiv Gandhi Univ. of Health Sciences, 1999)",
    "MBBS — Bangalore Medical College",
    "Medical Council: KMC 37,292 • DNB / NBE examiner • NABH Internal Auditor",
    "25,000+ clinical retinal cases: IRDs, ROP, DR, RVO, AMD, macular edema",
    "Languages: English, Hindi, Konkani, Kannada, Telugu, Tamil",
]:
    doc.add_paragraph(item, style="List Bullet")

# === PUBLICATIONS MOST RELEVANT ===
doc.add_heading("PUBLICATIONS — IRD & GENETICS LED", level=2)
for item in [
    "Inherited retinal disorders — genotype–phenotype correlation in an Indian cohort and "
    "the importance of genetic testing and counseling. Graefe's Archive, January 2023.",
    "Next generation sequencing and genetic counseling in ocular disorders — clinical, genetic, "
    "psychological and cultural dimensions. Indian Journal of Ophthalmology, 2026.",
    "Tele-ophthalmology and preventable childhood blindness: KIDROP experience. 2022.",
    "Clinically undetected macular changes in early ROP on SD-OCT. IOVS, 2011.",
    "Influence of foveal photoreceptor sub-elements on visual acuity in premature infants. IOVS, 2012.",
    "Outcomes of protocol-based management for Zone-I ROP. AJO, 2011.",
    "Spectral-domain OCT — Limitations and advances. Opening chapter, Elsevier \"Textbook and "
    "Atlas on OCT,\" 2013.",
]:
    doc.add_paragraph(item, style="List Bullet")
p = doc.add_paragraph()
add_hyperlink(p, f"{BASE_URL}/publications.html", "View all peer-reviewed publications")

# === CONFERENCE PRESENTATIONS ===
doc.add_heading("CONFERENCE PRESENTATIONS (Selected)", level=2)
for item in [
    "Speaker — Electrophysiology and Applied Clinical Genetics, Bangalore (Feb 2023)",
    "Homozygosity by descent of a foveal hypoplasia gene in an inbred South Indian family — AIOS Delhi (2015)",
    "Juvenile association of nephronophthisis with RP — AIOS Coimbatore (2017); APVRS Seoul (2018)",
    "Effect of consanguinity on retinitis pigmentosa severity — Asia ARVO Singapore (2011)",
    "BEST's vitelliform dystrophy with CNVM on SD-OCT — APAO Hyderabad (2013)",
    "\"Creating a safety net for ROP in India\" — AAO Retina Subspecialty, Chicago (2010)",
    "Speaker — Retina Summit, Goa (Feb 2024); Faculty — Retina Clinix International (2024–2025)",
    "Instructor — SSTC / DSTC courses, AIOC New Delhi (2025); AIOC Jaipur (2026)",
]:
    doc.add_paragraph(item, style="List Bullet")
p = doc.add_paragraph()
add_hyperlink(p, f"{BASE_URL}/presentations.html", "View all 120+ conference presentations")

# === LEADERSHIP, EDUCATION, NETWORK ===
doc.add_heading("LEADERSHIP, EDUCATION & NETWORK", level=2)
doc.add_heading("Patient Pathways & Treatment Centers", level=3)
for item in [
    "Multi-center retinal screening across 50+ sites and 10,000+ patient records (KIDROP)",
    "Referral workflow design, clinical protocols, treatment-center activation",
    "Imaging-based patient identification and longitudinal monitoring (OCT, OCTA, FAF)",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Physician Education & KOL", level=3)
for item in [
    "DNB / NBE examiner (clinical excellence standards in ophthalmology)",
    "16+ supervised DNB theses concentrated in retinal imaging and retinal disease",
    "120+ conference talks; faculty for AAO India Chapter, Retina Summit, Retina Clinix",
    "Hospital Infection Control committee leadership; NABH-aligned clinical governance",
]:
    doc.add_paragraph(item, style="List Bullet")

# === THESIS SUPERVISION ===
doc.add_heading("THESIS SUPERVISION (Selected)", level=2)
for item in [
    "Dexamethasone implant effect on macular perfusion (OCTA) in RVO — DNB (2021)",
    "Retinal microvasculature in stage-5 CKD on dialysis using SS-OCTA — DNB (2021)",
    "Longitudinal foveal morphology on SD-OCT in preterm infants with/without ROP — DNB (2012)",
    "Retinal vessel diameter in OCTA in hypertensive retinopathy — DNB (2024)",
    "OCTA changes in sleep apnea syndrome — DNB (2025)",
]:
    doc.add_paragraph(item, style="List Bullet")
p = doc.add_paragraph()
add_hyperlink(p, f"{BASE_URL}/theses.html", "View all 16+ theses supervised")

# === AWARDS ===
doc.add_heading("AWARDS & RECOGNITION", level=2)
for item in [
    "Gold Medal in MS Ophthalmology (Rajiv Gandhi Univ., 1999)",
    "Innovative Healthcare through PPP — KIDROP (2012)",
    "Best Pediatric Ophthalmology paper — APROP study",
    "Best ROP paper — BOS (2022); H.J. Mehta Award (best undergraduate research)",
    "11+ awards for research, innovation, and clinical excellence",
]:
    doc.add_paragraph(item, style="List Bullet")

# === FOOTER ===
p = doc.add_paragraph()
p.add_run("For comprehensive details: ")
add_hyperlink(p, f"{BASE_URL}/", "Full Online Resume (Ocular Gene Therapy & IRD version)", bold=True)

p = doc.add_paragraph()
p.add_run("drsherine@gmail.com | +91 98453 66008 | ")
add_hyperlink(p, f"{BASE_URL}/index.html#contact", "Contact")

doc.save("spark-resume.docx")
print("Wrote spark-resume.docx")
