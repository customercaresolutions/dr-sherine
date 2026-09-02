"""Shared layout helpers for the generated resumes.

All three resumes (clinical, corporate, ocular gene therapy) share one page setup and
one typographic scale, tuned so a full resume fits two pages: 8.5pt bullets, 9pt body,
10pt section headings, and per-paragraph spacing in the 0.5-6pt range.

The Word document defaults inherited from the template are 11pt type, 10pt space after
every paragraph and 1.15 line spacing — roughly 50% taller. Every paragraph these
helpers emit therefore sets its size and spacing explicitly; leaving either unset
silently falls back to those defaults and pushes the document onto a third page.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

TEMPLATE = "resume.docx"
HYPERLINK_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

BODY_SIZE = 9
BULLET_SIZE = 8.5
HEADING_SIZE = 10
SMALL_SIZE = 8


def new_document(template=TEMPLATE):
    """Clone the template for its margins and styles, then empty the body."""
    doc = Document(template)

    for el in list(doc.element.body):
        if el.tag == qn("w:p") or el.tag == qn("w:tbl"):
            doc.element.body.remove(el)

    # Drop stale hyperlink relationships inherited from the template
    for rel_id in [r for r, rel in doc.part.rels.items() if rel.reltype == HYPERLINK_TYPE]:
        del doc.part.rels[rel_id]

    return doc


def add_hyperlink(paragraph, url, text, bold=False):
    """Append a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, HYPERLINK_TYPE, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    if bold:
        rPr.append(OxmlElement("w:b"))
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def para(doc, text=None, style=None, size=None, bold=False, after=None, before=None, center=False):
    """Add a paragraph with explicit size and spacing."""
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    return p


def heading(doc, text, after=2):
    """Section heading. The run size is explicit because the Heading 2 style is 13pt,
    which overflows the two-page layout."""
    p = doc.add_heading("", level=2)
    p.add_run(text).font.size = Pt(HEADING_SIZE)
    p.paragraph_format.space_after = Pt(after)
    return p


def subheading(doc, text):
    return doc.add_heading(text, level=3)


def role(doc, title, dates, description):
    """Experience entry: bold role, lighter date range, then a description on a new line."""
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(BODY_SIZE)
    r = p.add_run(f"  ({dates})")
    r.font.size = Pt(SMALL_SIZE)
    r = p.add_run()
    r.add_break()
    r.add_text(description)
    r.font.size = Pt(BULLET_SIZE)
    p.paragraph_format.space_after = Pt(1)
    return p


def bullets(doc, items, after=1.0, size=BULLET_SIZE):
    for item in items:
        para(doc, item, style="List Bullet", size=size, after=after)


def header_block(doc, name, tagline, contact, link_url, link_text, name_size=12, link_bold=False):
    para(doc, name, size=name_size, bold=True, after=0, center=True)
    para(doc, tagline, size=BODY_SIZE, after=2, center=True)
    para(doc, contact, size=SMALL_SIZE, after=6, center=True)
    p = para(doc, after=6, center=True)
    add_hyperlink(p, link_url, link_text, bold=link_bold)


def link_para(doc, url, text, after=3):
    p = para(doc, after=after)
    add_hyperlink(p, url, text)
    return p


def footer_block(doc, link_url, link_text, contact, link_bold=False, contact_url=None):
    p = para(doc, size=SMALL_SIZE, after=2, before=6, center=True)
    p.add_run("For comprehensive details: ").font.size = Pt(SMALL_SIZE)
    add_hyperlink(p, link_url, link_text, bold=link_bold)
    p = para(doc, size=SMALL_SIZE, after=0, center=True)
    run = p.add_run(f"{contact} | " if contact_url else contact)
    run.font.size = Pt(SMALL_SIZE)
    if contact_url:
        add_hyperlink(p, contact_url, "Contact")
