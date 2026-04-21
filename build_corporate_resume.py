"""Generate corporate-resume.docx — 2-page corporate-positioning resume with links to the corporate site."""

from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, Emu

BASE_URL = "https://customercaresolutions.github.io/dr-sherine/versions/corporate"


def add_hyperlink(paragraph, url, text, bold=False):
    """Append a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # style as Hyperlink
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


doc = Document("resume.docx")  # clone margins/styles from existing resume
# Clear all existing content
for el in list(doc.element.body):
    if el.tag == qn("w:p") or el.tag == qn("w:tbl"):
        doc.element.body.remove(el)

# Drop stale hyperlink relationships inherited from the template
HYPERLINK_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
for rel_id in [r for r, rel in doc.part.rels.items() if rel.reltype == HYPERLINK_TYPE]:
    del doc.part.rels[rel_id]

# Re-add a trailing sectPr (needed to preserve page setup)
# (cloned from template — python-docx keeps the sectPr intact since we only removed p/tbl)

# === HEADER ===
p = doc.add_paragraph()
run = p.add_run("DR. SHERINE MARINA D'SOUZA BRAGANZA")
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph(
    "Clinical Product Strategist | Diagnostic Platform Development | Retinal Innovation"
)

p = doc.add_paragraph("drsherine@gmail.com | +91 98453 66008 | MS Ophthalmology (Gold Medal) | Bangalore, India")

p = doc.add_paragraph()
add_hyperlink(p, f"{BASE_URL}/", "Full Online Corporate Resume", bold=True)

# === PROFESSIONAL SUMMARY ===
doc.add_heading("PROFESSIONAL SUMMARY", level=2)
doc.add_paragraph(
    "Clinical Product Strategist with 25+ years in retinal disease diagnosis and clinical validation. "
    "Developed diagnostic platforms serving 50+ centers and 10,000+ patients. Expert in clinical "
    "validation frameworks, product lifecycle management, KOL engagement, regulatory strategy "
    "(FDA, CE Mark, NMPA), and clinical evidence generation. Proven ability to translate complex "
    "clinical problems into scalable diagnostic solutions and guide cross-functional product teams "
    "through validation and market adoption."
)

# === CORE EXPERTISE ===
doc.add_heading("CORE EXPERTISE", level=2)
for item in [
    "Diagnostic Platform Strategy • Product Lifecycle Management • Clinical Market Adoption",
    "Clinical Validation Frameworks • Diagnostic Accuracy Assessment • IRB/Ethics Protocol Design",
    "Regulatory Strategy (FDA, CE Mark, NMPA) • Clinical Evidence Generation • Reimbursement",
    "KOL Management & Clinical Advisory Boards • Cross-Functional Product Leadership",
    "Oculomics & Biomarker Research • OCT/OCTA Imaging • Retinal Disease Diagnostics",
    "Healthcare Systems Integration • EMR/EHR Design • NABH, HIPAA, Clinical Audit Standards",
]:
    doc.add_paragraph(item, style="List Bullet")

# === PROFESSIONAL EXPERIENCE ===
doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)

doc.add_paragraph(
    "Narayana Nethralaya II — Senior Consultant, Clinical Product Strategy  (Nov 2006 – Present)\n"
    "Lead clinical strategy and product development for diagnostic platforms in retinal disease "
    "detection. Guide platform lifecycle from problem identification through validation, regulatory "
    "strategy, and market adoption. Build clinical advisory boards and manage KOL engagement."
)

doc.add_paragraph(
    "KIDROP Initiative — Chief Clinical Officer & Platform Lead  (2015 – Present)\n"
    "Architected and scaled a diagnostic platform serving 50+ primary care centers and 10,000+ "
    "patients. Designed clinical workflows, diagnostic protocols, and algorithm requirements for "
    "ROP risk stratification. Authored clinical evidence publications supporting reimbursement."
)

doc.add_paragraph(
    "Lion's Eye Hospital — Clinical Operations & Research Lead  (Dec 2003 – Nov 2006)\n"
    "Established retinal imaging protocols and led clinical research initiatives. Collaborated on "
    "EMR design for research data collection and clinical workflows."
)

doc.add_paragraph(
    "Prathima Institute of Medical Sciences — Assistant Professor  (Feb 2002 – Present)\n"
    "Teaching clinical methodology, research design, and emerging diagnostic technologies."
)

p = doc.add_paragraph()
add_hyperlink(p, f"{BASE_URL}/experience.html", "View full professional experience")

# === EDUCATION & CREDENTIALS ===
doc.add_heading("EDUCATION & CREDENTIALS", level=2)
for item in [
    "MS Ophthalmology — Gold Medal (Rajiv Gandhi Univ. of Health Sciences, 1999)",
    "DNB Ophthalmology • Medical Council: KMC 37,292 • NABH Internal Auditor",
    "25,000+ clinical cases: ROP, diabetic retinopathy, RVO, macular edema",
    "Languages: English, Hindi, Konkani, Kannada, Telugu, Tamil",
]:
    doc.add_paragraph(item, style="List Bullet")

# === PUBLICATIONS & RESEARCH ===
doc.add_heading("PUBLICATIONS & RESEARCH", level=2)
doc.add_paragraph(
    "15+ peer-reviewed publications in IOVS, Ophthalmology, AJO, RETINA, Graefe's Archive, and IJO. "
    "Focus areas: diagnostic platform impact (KIDROP), clinical validation, imaging biomarkers "
    "(OCTA in DR/RVO/AMD/CKD), healthcare systems (REDROP), and genomic data integration."
)
p = doc.add_paragraph()
add_hyperlink(p, f"{BASE_URL}/publications.html", "View all peer-reviewed publications")

# === CONFERENCE PRESENTATIONS ===
doc.add_heading("CONFERENCE PRESENTATIONS (Selected)", level=2)
for item in [
    "AAO Chicago (2010): Safety net for ROP in India — platform scaling & clinical outcomes",
    "WOC Berlin (2010): Tele-ROP — The India Experience (diagnostic infrastructure)",
    "APAO (2021): OCTA microvascular changes in RVO — biomarker identification",
    "Retina Summit Goa (2024): Speaker — advances in retinal imaging & diagnostic innovation",
    "AIOS New Delhi (2025): Hodgkin's lymphoma causing CRAO — diagnostic case study",
    "AIOC Jaipur (2026): Instructor — SSTC/DSTC courses",
]:
    doc.add_paragraph(item, style="List Bullet")
p = doc.add_paragraph()
add_hyperlink(p, f"{BASE_URL}/presentations.html", "View all 120+ conference presentations")

# === LEADERSHIP & PLATFORM INNOVATION ===
doc.add_heading("LEADERSHIP & DIAGNOSTIC PLATFORM INNOVATION", level=2)
doc.add_heading("Platform Development & Clinical Validation", level=3)
for item in [
    "End-to-end diagnostic platform development — from clinical problem to market launch",
    "Clinical validation study design (IRB/ethics, diagnostic accuracy, clinical utility)",
    "Regulatory pathway strategy (FDA, CE Mark, NMPA) for diagnostic devices",
    "Clinical evidence generation for market launch and reimbursement",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("KOL, Advisory & Healthcare Systems", level=3)
for item in [
    "KOL network development and clinical advisory board management",
    "Healthcare compliance (NABH 5th edition auditor; HIPAA, clinical audit standards)",
    "Secretary, Hospital Infection Control committee",
    "DNB/NBE examiner; fellow mentoring; registered for AI in Medical Education (NBE, Jan 2026)",
]:
    doc.add_paragraph(item, style="List Bullet")

# === AWARDS & RECOGNITION ===
doc.add_heading("AWARDS & RECOGNITION", level=2)
for item in [
    "Gold Medal in MS Ophthalmology (Rajiv Gandhi Univ., 1999)",
    "Innovative Healthcare through PPP — KIDROP (2012)",
    "Best Pediatric Ophthalmology paper — APROP study",
    "Best ROP paper — BOS (2022); H.J. Mehta Award for best research work",
    "11+ awards for research, innovation, and clinical excellence",
]:
    doc.add_paragraph(item, style="List Bullet")

# === THESIS SUPERVISION ===
doc.add_heading("THESIS SUPERVISION (Selected)", level=2)
for item in [
    "Dexamethasone implant effect on macular perfusion (OCTA) in RVO — DNB (2021)",
    "Retinal microvasculature abnormalities in CKD on dialysis — DNB (2021)",
    "ROP trends: rural vs. urban in Karnataka — DNB (2021)",
    "OCTA changes in sleep apnea syndrome — DNB (2025)",
    "Retinal vessel diameter in OCTA and stereoscopic imaging in hypertensive retinopathy — DNB (2024)",
]:
    doc.add_paragraph(item, style="List Bullet")
p = doc.add_paragraph()
add_hyperlink(p, f"{BASE_URL}/theses.html", "View all 16 theses supervised")

# === FOOTER ===
p = doc.add_paragraph()
p.add_run("For comprehensive details: ")
add_hyperlink(p, f"{BASE_URL}/", "Full Online Corporate Resume", bold=True)

p = doc.add_paragraph()
p.add_run("drsherine@gmail.com | +91 98453 66008 | ")
add_hyperlink(p, f"{BASE_URL}/contact.html", "Contact")

doc.save("corporate-resume.docx")
print("Wrote corporate-resume.docx")
